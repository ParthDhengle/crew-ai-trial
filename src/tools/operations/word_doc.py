import os
import sys
import tempfile
from datetime import datetime
import warnings
import json
import re
from dotenv import load_dotenv
import platform
import subprocess

# Suppress syntax warnings
warnings.filterwarnings("ignore", category=SyntaxWarning)

# Absolute imports to avoid relative import issues
try:
    from utils.logger import setup_logger
    from common_functions.Find_project_root import find_project_root
except ImportError:
    def setup_logger():
        import logging
        logger = logging.getLogger("AIAssistant")
        logger.setLevel(logging.DEBUG)
        if not logger.handlers:
            console_handler = logging.StreamHandler()
            console_handler.setLevel(logging.INFO)
            formatter = logging.Formatter(
                "%(asctime)s - %(name)s - %(levelname)s - %(message)s",
                datefmt="%Y-%m-%d %H:%M:%S"
            )
            console_handler.setFormatter(formatter)
            logger.addHandler(console_handler)
        return logger
    
    def find_project_root():
        import os
        current_dir = os.path.abspath(os.path.dirname(__file__))
        root_dir = current_dir
        while root_dir != os.path.dirname(root_dir):
            if os.path.exists(os.path.join(root_dir, '.env')):  # Use '.git', 'requirements.txt', or another root marker if preferred
                return root_dir
            root_dir = os.path.dirname(root_dir)
        return current_dir  # Fallback to script's directory if no marker found

logger = setup_logger()
PROJECT_ROOT = find_project_root()

# Load .env file dynamically from project root
env_path = os.path.join(PROJECT_ROOT, '.env')
if not os.path.exists(env_path):
    logger.error(f".env file not found at {env_path}")
    sys.exit(1)
load_dotenv(env_path)
logger.info(f"Loaded .env file from {env_path}")

def check_word_installation():
    """Check if Microsoft Word is installed on the system."""
    try:
        if platform.system() == "Windows":
            # Common Word installation paths
            possible_paths = [
                r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE",
                r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE",
            ]
            
            for path in possible_paths:
                if os.path.exists(path):
                    logger.info(f"Microsoft Word found at: {path}")
                    return True, path
            
            logger.warning("Microsoft Word not found in common installation paths")
            return False, None
        else:
            logger.warning("Microsoft Word auto-open is only supported on Windows")
            return False, None
    except Exception as e:
        logger.error(f"Error checking Word installation: {str(e)}")
        return False, None

def call_grok(prompt: str, api_key: str) -> str:
    """Call Grok API directly."""
    try:
        from groq import Groq
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        logger.warning(f"Grok API failed: {str(e)}")
        return None

def call_gemini(prompt: str, api_key: str) -> str:
    """Call Gemini API as fallback."""
    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash-latest")
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        logger.warning(f"Gemini API failed: {str(e)}")
        return None

def clean_llm_response(response: str) -> str:
    """Clean LLM response to extract JSON."""
    if not response:
        return None
    
    logger.debug(f"Raw LLM response: {response[:500]}...")
    
    # Remove markdown code blocks
    response = re.sub(r'^```json\s*|\s*```$', '', response.strip(), flags=re.MULTILINE)
    response = re.sub(r'^```\s*|\s*```$', '', response.strip(), flags=re.MULTILINE)
    response = response.strip()
    
    # Try to find JSON within the response
    json_match = re.search(r'\{.*\}', response, re.DOTALL)
    if json_match:
        json_str = json_match.group(0)
        logger.debug(f"Extracted JSON: {json_str[:200]}...")
        return json_str
    
    logger.warning(f"No JSON found in response: {response[:200]}...")
    return None

def create_fallback_document_content(query: str) -> dict:
    """Create fallback document content when LLM fails."""
    logger.info("Creating fallback document content")
    
    return {
        "title": f"Report for: {query}",
        "sections": [
            {
                "heading": "Introduction",
                "content": "This is a basic report generated as fallback. Please provide a more detailed query for better results."
            },
            {
                "heading": "Summary",
                "content": "No specific content generated due to LLM failure."
            }
        ]
    }

def create_word_document(output_path: str, content: dict):
    """Create a Microsoft Word document using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Add title
        doc.add_heading(content.get("title", "Generated Report"), 0)
        
        # Add sections
        for section in content.get("sections", []):
            doc.add_heading(section.get("heading", "Section"), level=1)
            doc.add_paragraph(section.get("content", ""))
        
        doc.save(output_path)
        logger.info(f"Word document created at: {output_path}")
        return True
    except ImportError:
        logger.error("python-docx library not found. Please install it with 'pip install python-docx'")
        return False
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        return False

def word_generate_from_query(query: str) -> tuple[bool, str]:
    """
    Generates a Microsoft Word document based on a user query using AI-driven content generation.
    
    Args:
        query (str): User query describing desired report (e.g., "Write a report on climate change").
    
    Returns:
        tuple[bool, str]: (success, result message or error)
    """
    try:
        # Check Word installation
        word_installed, word_path = check_word_installation()
        if not word_installed:
            logger.warning("Microsoft Word not found. Document will be created but may not open automatically.")
        
        # Verify API keys
        grok_api_key = os.getenv("GROQ_API_KEY3")
        gemini_api_key = os.getenv("GEMINI_API_KEY3")
        if not grok_api_key and not gemini_api_key:
            logger.error(f"No API keys found for GROQ_API_KEY1 or GEMINI_API_KEY1 in {env_path}")
            return False, f"Error: No API keys found for GROQ_API_KEY1 or GEMINI_API_KEY1 in {env_path}"

        logger.info(f"Generating Word document for query: {query}")
        
        # Construct improved prompt for LLM
        prompt = f"""
You are a report writing expert. Generate content for a Word document based on this request.

User Query: "{query}"

Return ONLY a valid JSON object with this exact structure:
{{
  "title": "Report Title",
  "sections": [
    {{
      "heading": "Section Heading",
      "content": "Section content text..."
    }}
  ]
}}

Make the content informative and relevant to the query.
Generate summaries, insights, or descriptions based on the query.
Return only the JSON, no explanation.
"""
        
        # Call LLM with retry logic
        response = None
        content = None
        
        for attempt in range(3):
            logger.info(f"Attempting LLM call #{attempt + 1}")
            
            if grok_api_key and not response:
                response = call_grok(prompt, grok_api_key)
                if response:
                    cleaned = clean_llm_response(response)
                    if cleaned:
                        try:
                            content = json.loads(cleaned)
                            logger.info("Successfully parsed Grok response")
                            break
                        except json.JSONDecodeError as e:
                            logger.warning(f"Grok JSON decode error: {str(e)}")
                            response = None
            
            if gemini_api_key and not content:
                response = call_gemini(prompt, gemini_api_key)
                if response:
                    cleaned = clean_llm_response(response)
                    if cleaned:
                        try:
                            content = json.loads(cleaned)
                            logger.info("Successfully parsed Gemini response")
                            break
                        except json.JSONDecodeError as e:
                            logger.warning(f"Gemini JSON decode error: {str(e)}")
                            response = None
        
        # Use fallback if LLM failed
        if not content:
            logger.warning("LLM parsing failed after all retries. Using fallback content.")
            content = create_fallback_document_content(query)
        
        logger.info(f"Document content: {content}")
        
        # Create temporary file for document
        output_dir = tempfile.mkdtemp(prefix="word_document_")
        document_name = f"auto_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        document_path = os.path.join(output_dir, document_name)
        
        if not create_word_document(document_path, content):
            return False, "Failed to create Word document. Ensure python-docx is installed."
        
        success_message = f"Word document created at: {document_path}"
        
        if word_installed:
            try:
                # Try to open with Word if available
                os.startfile(document_path)
                success_message += "\nDocument opened in Microsoft Word."
            except Exception as e:
                logger.warning(f"Failed to open document: {str(e)}")
                success_message += f"\nCould not auto-open document: {str(e)}"
        else:
            # Open the directory instead if Word not found
            if platform.system() == "Windows":
                os.startfile(output_dir)
            success_message += "\nDocument directory opened."
        
        logger.info(success_message)
        return True, success_message
        
    except Exception as e:
        logger.error(f"Error generating Word document: {str(e)}")
        return False, f"Error: {str(e)}"

if __name__ == "__main__":
    # Ensure environment variables are loaded
    grok_api_key = os.getenv("GROQ_API_KEY3")
    gemini_api_key = os.getenv("GEMINI_API_KEY3")
    if not grok_api_key and not gemini_api_key:
        print(f"Error: Please set GROQ_API_KEY1 or GEMINI_API_KEY1 in {env_path}")
        sys.exit(1)
    
    # Example usage
    test_query = "Write a detailed report on Mahabharat"
    success, result = word_generate_from_query(test_query)
    print(f"Success: {success}")
    print(result)